from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submission" / "TE-KG_working_draft_v0.1.docx"
OUTPUT = ROOT / "submission" / "TE-KG_working_draft_v0.1.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0x9B, 0x1C, 0x1C)
LIGHT_GRAY = "E9ECEF"


def get_style(doc, name):
    """Return a style by its displayed name, including Pandoc heading styles."""
    return next(style for style in doc.styles if style.name == name)


def set_run_font(run, name="Times New Roman", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, size, bold=False, color=BLACK, before=0, after=6, line=1.15):
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.extend([tbl_w, tbl_ind, layout])

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 90), ("bottom", 90), ("start", 120), ("end", 120)):
                elem = margins.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    margins.append(elem)
                elem.set(qn("w:w"), str(value))
                elem.set(qn("w:type"), "dxa")


def insert_before(block, paragraph):
    block.addprevious(paragraph._p)


def build():
    doc = Document(SOURCE)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    set_style(get_style(doc, "Normal"), 11, before=0, after=6, line=1.15)
    for name in ("Body Text", "First Paragraph"):
        set_style(get_style(doc, name), 11, before=0, after=6, line=1.15)
    set_style(get_style(doc, "Title"), 18, bold=True, before=0, after=8, line=1.0)
    set_style(get_style(doc, "Author"), 11, before=0, after=2, line=1.0)
    set_style(get_style(doc, "Date"), 10, color=MUTED, before=0, after=8, line=1.0)
    set_style(get_style(doc, "Abstract"), 10.5, before=0, after=8, line=1.1)
    set_style(get_style(doc, "Heading 1"), 14, bold=True, before=14, after=6, line=1.0)
    set_style(get_style(doc, "Heading 2"), 12, bold=True, before=10, after=4, line=1.0)
    set_style(get_style(doc, "Bibliography"), 9, before=0, after=3, line=1.0)
    set_style(get_style(doc, "Caption"), 9.5, bold=False, before=6, after=4, line=1.0)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name in {"Heading 1", "Heading 2"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name in {"Title", "Author", "Date"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style.name == "Bibliography":
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        for run in paragraph.runs:
            if run.text.startswith("[Draft note:"):
                run.font.color.rgb = RED
                run.italic = True

    title = doc.paragraphs[0]
    for run in title.runs:
        set_run_font(run, size=18, bold=True)

    abstract = next(p for p in doc.paragraphs if p.style.name == "Abstract")
    abstract_heading = doc.add_paragraph("Abstract")
    abstract_heading.style = get_style(doc, "Heading 1")
    abstract._p.addprevious(abstract_heading._p)

    date = next(p for p in doc.paragraphs if p.style.name == "Date")
    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_after = Pt(8)
    warning_run = warning.add_run("WORKING DRAFT - unresolved author and provenance fields are shown in red")
    set_run_font(warning_run, size=9, italic=True, color=RED)
    date._p.addnext(warning._p)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Keywords:") or paragraph.text.startswith("Database URL:"):
            if paragraph.runs:
                paragraph.runs[0].bold = True
        if "values were transformed as , and TE-gene" in paragraph.text:
            for run in paragraph.runs:
                if "values were transformed as , and TE-gene" in run.text:
                    run.text = run.text.replace(
                        "values were transformed as , and TE-gene",
                        "values were transformed as log2(count + 1), and TE-gene",
                    )

    bibliography = [p for p in doc.paragraphs if p.style.name == "Bibliography"]
    if bibliography:
        heading = doc.add_paragraph("References")
        heading.style = get_style(doc, "Heading 1")
        bibliography[0]._p.addprevious(heading._p)
        dois = [
            "10.1186/s13100-015-0041-9",
            "10.1186/s13100-017-0107-y",
            "10.1093/nar/gks1265",
            "10.1093/nar/30.1.205",
            "10.1002/humu.20307",
            "10.1093/nar/gku1043",
            "10.1093/nar/gkm949",
            "10.1093/nar/gkad904",
            "10.1128/jvi.00059-23",
            "10.1093/nar/gkaf1235",
            "10.1369/0022155414562646",
            "10.1126/science.1260419",
            "10.1038/s41586-019-1186-3",
            "10.1111/j.2517-6161.1995.tb02031.x",
            "10.1088/1742-5468/2008/10/P10008",
            "10.1038/s41576-020-0251-y",
        ]
        for paragraph, doi in zip(bibliography, dois):
            paragraph.add_run(" ")
            add_hyperlink(paragraph, f"doi:{doi}", f"https://doi.org/{doi}")

    if doc.tables:
        table = doc.tables[0]
        if table.rows and not any(cell.text.strip() for cell in table.rows[-1].cells):
            table._tbl.remove(table.rows[-1]._tr)
        set_table_geometry(table, [1900, 850, 2750, 3860])
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                tr_pr = row._tr.get_or_add_trPr()
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    shade_cell(cell, LIGHT_GRAY)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    for run in paragraph.runs:
                        set_run_font(run, size=8.5, bold=(row_idx == 0))
        caption = doc.add_paragraph(
            "Table 1. Current TE-KG content snapshot. Counts describe different runtime units and must not be summed. The snapshot was verified on 31 July 2026."
        )
        caption.style = get_style(doc, "Caption")
        table._tbl.addprevious(caption._p)

    header = section.header.paragraphs[0]
    header.text = "TE-KG | Working draft v0.1"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer, " PAGE ")

    props = doc.core_properties
    props.title = "TE-KG working manuscript draft v0.1"
    props.subject = "Database journal manuscript working draft"
    props.author = "Authors to be supplied"
    props.keywords = "TE-KG; transposable elements; knowledge graph; database"
    props.comments = "Generated from the canonical LaTeX working draft; Author Contributions deferred."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
