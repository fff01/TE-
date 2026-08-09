from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ARTICLE_ROOT = Path(__file__).resolve().parents[1]
DRAFTS = ARTICLE_ROOT / "drafts"
SUBMISSION = ARTICLE_ROOT / "submission"
SOURCE_EN = SUBMISSION / "TE-KG_working_draft_v0.1.docx"
SOURCE_ZH = DRAFTS / "word_source_zh.tex"
OUTPUT_EN = SUBMISSION / "TE-KG_Database_submission_draft_en.docx"
OUTPUT_ZH = SUBMISSION / "TE-KG_Database_submission_draft_zh.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
RED = RGBColor(0x9B, 0x1C, 0x1C)
BLUE = "0563C1"
LIGHT_GRAY = "E9ECEF"

DOIS = [
    "10.1186/s13100-015-0041-9",
    "10.1186/s13100-017-0107-y",
    "10.1093/nar/gks1265",
    "10.1093/nar/30.1.205",
    "10.1002/humu.20307",
    "10.1093/nar/gku1043",
    "10.1093/nar/gkm949",
    "10.1093/nar/gkad904",
    "10.1128/jvi.00059-23",
    "10.1093/nar/gkaf1235",
    "10.1369/0022155414562646",
    "10.1126/science.1260419",
    "10.1038/s41586-019-1186-3",
    "10.1111/j.2517-6161.1995.tb02031.x",
    "10.1088/1742-5468/2008/10/P10008",
    "10.1038/s41576-020-0251-y",
]

ZH_TABLE_ROWS = [
    ("组成部分及单位", "数量", "定义", "来源与边界"),
    ("Neo4j TE 节点", "225", "tekg3 图中的 TE 实体", "实时标签计数查询；与 Browse 条目不是同一单位"),
    ("Neo4j Paper 节点", "2,308", "图中表示的保留论文", "实时标签计数查询；与最终文献语料库一致"),
    ("有向生物学关系", "12,444", "已存储的 BIO_RELATION 关系", "按存储方向计数一次；当前全部记录均含谓词和 PMID 来源字段"),
    ("生物医学实体类别", "11", "碳水化合物、疾病、功能、基因、脂质、突变、肽、药物、蛋白质、RNA 和毒素", "不包括 Paper、DiseaseCategory 和分类标签"),
    ("Browse 目录条目", "276", "版本化 TE 目录记录", "由 MySQL 支持的 Browse API；与图 TE 节点属于不同记录单位"),
    ("已分类 TE", "225 个中的 192 个", "实时摘要中已具有分类类别的 TE 节点", "由 Neo4j 支持的分类 API"),
    ("表达样本", "1,158", "205 个正常组织、307 个正常原代细胞和 646 个癌细胞系样本", "三类来源组；不是跨情境配对队列"),
    ("可检索共表达条目", "784", "285 个 TE 和 499 个基因条目", "三类情境的经批准展示目录；不是完整离线网络的特征数量"),
    ("下载文件", "10", "6 个表达文件、2 个图文件和 2 个分类文件", "当前网页入口；仍需公开归档发布"),
]


def get_style(doc: Document, name: str):
    return next(style for style in doc.styles if style.name == name)


def set_rfonts(element, western: str, east_asia: str) -> None:
    rpr = element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def style_definition(
    style,
    *,
    size: float,
    bold: bool = False,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 2.0,
    western: str = "Times New Roman",
    east_asia: str = "Times New Roman",
) -> None:
    style.font.name = western
    set_rfonts(style._element, western, east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def apply_run_font(run, *, size: float | None = None, western: str, east_asia: str) -> None:
    run.font.name = western
    set_rfonts(run._element, western, east_asia)
    if size is not None:
        run.font.size = Pt(size)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str, east_asia: str) -> None:
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr, value in (("ascii", "Times New Roman"), ("hAnsi", "Times New Roman"), ("eastAsia", east_asia)):
        fonts.set(qn(f"w:{attr}"), value)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    rpr.extend([fonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_line_numbering(section) -> None:
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:lnNumType"))
    if old is not None:
        sect_pr.remove(old)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "newPage")
    sect_pr.append(line_numbers)


def configure_sections(doc: Document, east_asia: str) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        set_line_numbering(section)

        for paragraph in section.header.paragraphs:
            clear_paragraph(paragraph)

        footer = section.footer.paragraphs[0]
        clear_paragraph(footer)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.space_after = Pt(0)
        footer.paragraph_format.line_spacing = 1.0
        run = footer.add_run("Page ")
        apply_run_font(run, size=9, western="Times New Roman", east_asia=east_asia)
        add_field(footer, " PAGE ")


def configure_styles(doc: Document, east_asia: str) -> None:
    for name in ("Normal", "Body Text", "First Paragraph", "Abstract"):
        style_definition(
            get_style(doc, name),
            size=12,
            before=0,
            after=0,
            line_spacing=2.0,
            east_asia=east_asia,
        )
    style_definition(get_style(doc, "Title"), size=16, bold=True, after=8, line_spacing=1.0, east_asia=east_asia)
    style_definition(get_style(doc, "Author"), size=12, after=3, line_spacing=1.0, east_asia=east_asia)
    style_definition(get_style(doc, "Date"), size=10, after=6, line_spacing=1.0, east_asia=east_asia)
    style_definition(get_style(doc, "Heading 1"), size=12, bold=True, before=12, after=6, line_spacing=1.0, east_asia=east_asia)
    style_definition(get_style(doc, "Heading 2"), size=12, bold=True, before=8, after=4, line_spacing=1.0, east_asia=east_asia)
    style_definition(get_style(doc, "Bibliography"), size=10, after=3, line_spacing=1.0, east_asia=east_asia)
    style_definition(get_style(doc, "Caption"), size=10, after=4, line_spacing=1.0, east_asia=east_asia)


def style_document_content(doc: Document, east_asia: str) -> None:
    centered = {"Title", "Author", "Date"}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        paragraph.paragraph_format.widow_control = True
        if style_name in {"Heading 1", "Heading 2", "Caption"}:
            paragraph.paragraph_format.keep_with_next = True
        if style_name in centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name in {"Normal", "Body Text", "First Paragraph", "Abstract"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Inches(0)
        if style_name == "Bibliography":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        for run in paragraph.runs:
            apply_run_font(run, western="Times New Roman", east_asia=east_asia)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.extend([tbl_w, tbl_ind, layout])

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 90), ("bottom", 90), ("start", 120), ("end", 120)):
                elem = margins.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    margins.append(elem)
                elem.set(qn("w:w"), str(value))
                elem.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "B7B7B7")


def style_table(table, east_asia: str) -> None:
    set_table_geometry(table, [1900, 850, 2750, 3860])
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = tr_pr.find(qn("w:tblHeader"))
            if repeat is None:
                repeat = OxmlElement("w:tblHeader")
                tr_pr.append(repeat)
            repeat.set(qn("w:val"), "true")
        for column_index, cell in enumerate(row.cells):
            if row_index == 0:
                shade_cell(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index == 1 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    apply_run_font(run, size=8.5, western="Times New Roman", east_asia=east_asia)
                    run.bold = row_index == 0


def add_warning(doc: Document, text: str, east_asia: str) -> None:
    if any(paragraph.text == text for paragraph in doc.paragraphs):
        return
    date = next(paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Date")
    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_after = Pt(8)
    warning.paragraph_format.line_spacing = 1.0
    run = warning.add_run(text)
    apply_run_font(run, size=9, western="Times New Roman", east_asia=east_asia)
    run.italic = True
    run.font.color.rgb = RED
    date._p.addnext(warning._p)


def ensure_heading_before(doc: Document, paragraph, text: str) -> None:
    previous = paragraph._p.getprevious()
    if previous is not None and "".join(previous.itertext()).strip() == text:
        return
    heading = doc.add_paragraph(text, style=get_style(doc, "Heading 1"))
    paragraph._p.addprevious(heading._p)


def mark_draft_notes(doc: Document, marker: str) -> None:
    for paragraph in doc.paragraphs:
        if marker not in paragraph.text:
            continue
        for run in paragraph.runs:
            if marker in run.text or run.font.color.rgb == RED:
                run.font.color.rgb = RED
                run.italic = True


def move_english_front_matter(doc: Document) -> None:
    database_url = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("Database URL:"))
    keywords = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("Keywords:"))
    keywords._p.addprevious(database_url._p)


def move_funding_before_acknowledgements(doc: Document, funding: str, acknowledgements: str) -> None:
    funding_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == funding)
    funding_body = funding_heading._p.getnext()
    acknowledgements_heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == acknowledgements)
    acknowledgements_heading._p.addprevious(funding_heading._p)
    acknowledgements_heading._p.addprevious(funding_body)


def replace_paragraph_text(paragraph, old: str, new: str, east_asia: str) -> None:
    if old not in paragraph.text:
        return
    replacement = paragraph.text.replace(old, new)
    clear_paragraph(paragraph)
    run = paragraph.add_run(replacement)
    apply_run_font(run, western="Times New Roman", east_asia=east_asia)


def configure_properties(doc: Document, language: str) -> None:
    props = doc.core_properties
    props.title = f"TE-KG Database submission working draft ({language})"
    props.subject = "Database journal submission manuscript working draft"
    props.author = "Authors to be supplied"
    props.keywords = "TE-KG; transposable elements; knowledge graph; Database"
    props.comments = "Reformatted for Database submission; unresolved author and provenance fields remain marked."


def build_english() -> None:
    doc = Document(SOURCE_EN)
    configure_sections(doc, "Times New Roman")
    configure_styles(doc, "Times New Roman")
    move_english_front_matter(doc)
    move_funding_before_acknowledgements(doc, "Funding", "Acknowledgements")
    browse = next(paragraph for paragraph in doc.paragraphs if "Browse and Search resolve TE names" in paragraph.text)
    replace_paragraph_text(
        browse,
        "Browse and Search resolve TE names and lead to structured identity records.",
        "Browse combines TE catalogue filtering, built-in search and access to structured records for selected TEs.",
        "Times New Roman",
    )
    style_document_content(doc, "Times New Roman")
    add_warning(doc, "WORKING DRAFT - unresolved author and provenance fields are shown in red", "Times New Roman")
    mark_draft_notes(doc, "[Draft note:")
    for table in doc.tables:
        style_table(table, "Times New Roman")
    configure_properties(doc, "English")
    doc.save(OUTPUT_EN)


def build_chinese_raw(target: Path) -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("Pandoc is required to build the Chinese manuscript.")
    subprocess.run(
        [
            pandoc,
            SOURCE_ZH.name,
            "--from=latex",
            "--to=docx",
            "--citeproc",
            "--csl=../submission/database.csl",
            "--bibliography=../references.bib",
            "--resource-path=.",
            f"--output={target}",
        ],
        cwd=DRAFTS,
        check=True,
    )


def remove_pandoc_table_residue(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if "组成部分及单位 & 数量 & 定义 & 来源与边界" in paragraph.text:
            paragraph._element.getparent().remove(paragraph._element)


def add_chinese_table(doc: Document) -> None:
    target = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "分类视图与目录视图保留不同的记录单位")
    caption = doc.add_paragraph(
        "表 1. 当前 TE-KG 内容快照。各项计数代表不同的运行时单位，不得相加。快照于 2026 年 7 月 31 日完成验证。",
        style=get_style(doc, "Caption"),
    )
    table = doc.add_table(rows=len(ZH_TABLE_ROWS), cols=4)
    for row, values in zip(table.rows, ZH_TABLE_ROWS):
        for cell, value in zip(row.cells, values):
            cell.text = value
    target._p.addprevious(caption._p)
    target._p.addprevious(table._tbl)
    style_table(table, "SimSun")


def repair_chinese_math(doc: Document) -> None:
    paragraph = next(paragraph for paragraph in doc.paragraphs if "丰度值按" in paragraph.text and "Spearman" in paragraph.text)
    replacement = (
        "共表达分别在正常组织、正常原代细胞和癌细胞系情境中计算。丰度值按 log2(count + 1) 转换，"
        "TE-基因关联使用 Spearman 秩相关衡量。候选边在 |r| ≥ 0.4 且 Benjamini–Hochberg "
        "错误发现率（FDR）不高于 0.05 时保留 (14)。保留的正相关边用于 Louvain 社区发现，"
        "随机种子为 42，分辨率为 1.8 (15)。系统将这些边解释为统计相关，而不是 TE 调控基因的证据。"
    )
    clear_paragraph(paragraph)
    run = paragraph.add_run(replacement)
    apply_run_font(run, western="Times New Roman", east_asia="SimSun")


def add_chinese_reference_heading_and_dois(doc: Document) -> None:
    bibliography = [paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Bibliography"]
    ensure_heading_before(doc, bibliography[0], "参考文献")
    for paragraph, doi in zip(bibliography, DOIS):
        if doi in paragraph.text:
            continue
        paragraph.add_run(" ")
        add_hyperlink(paragraph, f"doi:{doi}", f"https://doi.org/{doi}", "SimSun")


def build_chinese() -> None:
    with tempfile.TemporaryDirectory(prefix="tekg_database_zh_") as temp_dir:
        raw = Path(temp_dir) / "tekg_database_zh_raw.docx"
        build_chinese_raw(raw)
        doc = Document(raw)

    configure_sections(doc, "SimSun")
    configure_styles(doc, "SimSun")
    abstract = next(paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Abstract")
    ensure_heading_before(doc, abstract, "摘要")
    remove_pandoc_table_residue(doc)
    repair_chinese_math(doc)
    add_chinese_table(doc)
    add_chinese_reference_heading_and_dois(doc)
    style_document_content(doc, "SimSun")
    add_warning(doc, "中文审阅稿 - 尚未补齐的作者与来源信息以红色显示", "SimSun")
    mark_draft_notes(doc, "[待补说明：")
    configure_properties(doc, "Chinese")
    doc.save(OUTPUT_ZH)


def main() -> None:
    OUTPUT_EN.parent.mkdir(parents=True, exist_ok=True)
    build_english()
    build_chinese()
    print(OUTPUT_EN)
    print(OUTPUT_ZH)


if __name__ == "__main__":
    main()
